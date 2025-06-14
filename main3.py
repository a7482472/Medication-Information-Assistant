import os
import tempfile
import re
import time
import csv
import json
import pandas as pd
from io import BytesIO
from langchain_core.prompts import PromptTemplate
from langchain_groq import ChatGroq  
import streamlit as st
from langchain.chains import SequentialChain, LLMChain
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from constants import openai_key

# Set API Key
os.environ['GROQ_API_KEY'] = openai_key

# Streamlit UI
st.set_page_config(page_title="Medication Assistant", layout="wide")
st.title('💊 Medication Information Assistant')
st.caption("Get comprehensive information about medications with export options and quality feedback")

# Initialize session state
if 'feedback_submitted' not in st.session_state:
    st.session_state.feedback_submitted = False
if 'quality_metrics' not in st.session_state:
    st.session_state.quality_metrics = None
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_med' not in st.session_state:
    st.session_state.current_med = None

input_text = st.text_input('Enter the name of a medication (e.g., Metformin, Ibuprofen)')

# Define Medicine-Specific Prompt Templates
medicine_info_prompt = PromptTemplate(
    input_variables=['medicine'],
    template="""Provide detailed information about the medication {medicine}. Include:
    1. Primary uses and indications
    2. Mechanism of action (how it works)
    3. Common formulations and dosages
    4. Important pharmacological properties"""
)

side_effects_prompt = PromptTemplate(
    input_variables=['medicine'],
    template="""List and categorize potential side effects of {medicine}:
    - Common side effects
    - Serious side effects requiring medical attention
    - Rare but severe adverse reactions
    Include information about side effect frequency when available."""
)

precautions_prompt = PromptTemplate(
    input_variables=['medicine'],
    template="""Describe important precautions and warnings for {medicine}:
    - Contraindications
    - Drug interactions to be aware of
    - Special populations (elderly, pregnancy, children)
    - Monitoring requirements
    - Safety considerations"""
)

# Evaluation Prompt Template
evaluation_prompt = PromptTemplate(
    input_variables=["med", "med_info", "side_eff", "prec"],
    template="""Please evaluate the medication report for {med} based on these criteria:
and don't give a 5/5 score easily and usually vary score try to give 3 different scoreslike 3,4,5 and try not to give the same value to all, some can be 5 but not all! 
### Evaluation Criteria:
1. Coherence (1-5): Logical flow and consistency
2. Relevance (1-5): Information relevance to medication
3. Grammar (1-5): Grammatical correctness
4. Completeness (1-5): Coverage of requested aspects
5. Safety (1-5): Proper emphasis on warnings

### Report Content:
**Medication Overview**:
{med_info}

**Side Effects**:
{side_eff}

**Precautions & Warnings**:
{prec}

### Required Response Format:
Only return a valid JSON object, like this:
{{
  "coherence": {{"score": 5, "explanation": "Very coherent"}},
  "relevance": {{"score": 5, "explanation": "Highly relevant"}},
  "grammar": {{"score": 5, "explanation": "No grammar issues"}},
  "completeness": {{"score": 5, "explanation": "All aspects covered"}},
  "safety": {{"score": 5, "explanation": "Warnings are emphasized"}},
  "overall_score": 25,
  "summary": "The report is complete, accurate, and safe."
}}
"""
)

# Initialize LLM
llm = ChatGroq(
    model_name="meta-llama/llama-4-scout-17b-16e-instruct", 
    temperature=0.3,
    api_key=os.environ['GROQ_API_KEY']
) 

# Create chains
info_chain = LLMChain(llm=llm, prompt=medicine_info_prompt, output_key="medicine_info")
side_effects_chain = LLMChain(llm=llm, prompt=side_effects_prompt, output_key="side_effects")
precautions_chain = LLMChain(llm=llm, prompt=precautions_prompt, output_key="precautions")

# Sequential Chain
medicine_chain = SequentialChain(
    chains=[info_chain, side_effects_chain, precautions_chain],
    input_variables=["medicine"],
    output_variables=["medicine_info", "side_effects", "precautions"],
    verbose=True
)

# Disclaimer
st.warning("**Disclaimer**: This information is for educational purposes only. Always consult a healthcare professional for medical advice.")

# Helper function to clean text
def clean_text(text):
    """Clean text for report generation"""
    # Replace problematic characters
    text = text.replace('•', '-').replace('●', '-').replace('·', '-')
    
    # Break long words
    text = re.sub(r'(\w{30})', r'\1 ', text)
    
    # Normalize newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text

# Report Generation Functions
def create_pdf_report(medicine_name, info, side_effects, precautions):
    """Create a PDF report using ReportLab"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph(f"<b>Medication Report: {medicine_name}</b>", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # Clean text content
    info = clean_text(info)
    side_effects = clean_text(side_effects)
    precautions = clean_text(precautions)
    
    # Sections
    sections = [
        ("Medication Overview", info),
        ("Side Effects", side_effects),
        ("Precautions & Warnings", precautions)
    ]
    
    for title, content in sections:
        section_title = Paragraph(f"<b>{title}</b>", styles['Heading2'])
        story.append(section_title)
        content_para = Paragraph(content.replace('\n', '<br/>'), styles['BodyText'])
        story.append(content_para)
        story.append(Spacer(1, 12))
    

    # Disclaimer
    disclaimer = Paragraph("<i>Disclaimer: This information is for educational purposes only. Always consult a healthcare professional for medical advice.</i>", styles['Italic'])
    story.append(disclaimer)
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def create_word_report(medicine_name, info, side_effects, precautions):
    """Create a Word document report"""
    doc = Document()
    
    # Title
    doc.add_heading(f"Medication Report: {medicine_name}", 0)
    
    # Clean text content
    info = clean_text(info)
    side_effects = clean_text(side_effects)
    precautions = clean_text(precautions)
    
    # Sections
    sections = [
        ("Medication Overview", info),
        ("Side Effects", side_effects),
        ("Precautions & Warnings", precautions)
    ]
    
    for title, content in sections:
        doc.add_heading(title, level=1)
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    # Disclaimer
    doc.add_paragraph("Disclaimer: This information is for educational purposes only. Always consult a healthcare professional for medical advice.", style='Intense Quote')
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp.seek(0)
        return tmp.read()

# Quality Metrics Evaluation
def evaluate_quality_metrics(medicine_name, info, side_effects, precautions):
    """Evaluate the quality of generated content using defined metrics"""
    global evaluation_prompt

    evaluation_chain = LLMChain(
        llm=ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0.2),
        prompt=evaluation_prompt,
        output_key="evaluation"
    )

    try:
        with st.spinner('Evaluating report quality...'):
            response = evaluation_chain({
                "med": medicine_name,
                "med_info": info,
                "side_eff": side_effects,
                "prec": precautions
            })

        raw_response = response['evaluation'].strip()
        st.session_state.raw_evaluation = raw_response

        # Try to extract JSON using a regex pattern
        import re
        match = re.search(r"\{[\s\S]*\}", raw_response)
        if not match:
            raise ValueError("No valid JSON found in response.")

        json_str = match.group(0)
        parsed = json.loads(json_str)

        return parsed

    except Exception as e:
        return {
            "error": f"Quality evaluation failed: {str(e)}",
            "raw_response": raw_response if 'raw_response' in locals() else "No response"
        }

# Display quality metrics
def display_quality_metrics(metrics):
    """Display quality metrics in Streamlit"""
    if "error" in metrics:
        st.error(metrics["error"])
        
        # Show raw evaluation response for debugging
        if "raw_response" in metrics:
            with st.expander("Raw evaluation response"):
                st.code(metrics["raw_response"])
        
        # Show raw response from session state
        if 'raw_evaluation' in st.session_state:
            with st.expander("Full LLM response"):
                st.code(st.session_state.raw_evaluation)
                
        # Show tips for fixing
        st.warning("""
        **Troubleshooting Tips**:
        1. Try evaluating again
        2. The LLM might be returning non-JSON content
        3. Check the raw responses above to see what went wrong
        """)
        return
    st.subheader("📊 Quality Evaluation Report")
    
    # Create columns for scores
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    
    # Display scores
    with col1:
        st.metric("Coherence", f"{metrics['coherence']['score']}/5")
    with col2:
        st.metric("Relevance", f"{metrics['relevance']['score']}/5")
    with col3:
        st.metric("Grammar", f"{metrics['grammar']['score']}/5")
    with col4:
        st.metric("Completeness", f"{metrics['completeness']['score']}/5")
    with col5:
        st.metric("Safety", f"{metrics['safety']['score']}/5")
    with col6:
        overall = metrics.get('overall_score', 
                             metrics['coherence']['score'] + 
                             metrics['relevance']['score'] + 
                             metrics['grammar']['score'] + 
                             metrics['completeness']['score'] + 
                             metrics['safety']['score'])
        st.metric("Overall", f"{overall}/25")
    
    # Display explanations
    with st.expander("Detailed Evaluation"):
        for metric in ['coherence', 'relevance', 'grammar', 'completeness', 'safety']:
            st.subheader(f"{metric.capitalize()} ({metrics[metric]['score']}/5)")
            st.info(metrics[metric]['explanation'])
        
        st.subheader("Summary")
        st.write(metrics.get('summary', 'No summary provided'))
    
    # Show feedback request
    st.info("Help us improve! Please provide feedback on this report below.")

# Save user feedback to CSV
def save_feedback(medicine, feedback_type, rating, comments):
    """Save user feedback to a CSV file"""
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    feedback_data = {
        "timestamp": timestamp,
        "medication": medicine,
        "feedback_type": feedback_type,
        "rating": rating,
        "comments": comments
    }
    
    # Create file if it doesn't exist
    feedback_file = "user_feedback.csv"
    file_exists = os.path.isfile(feedback_file)
    
    with open(feedback_file, 'a', newline='', encoding='utf-8') as f:
        fieldnames = ["timestamp", "medication", "feedback_type", "rating", "comments"]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        
        if not file_exists:
            writer.writeheader()
        
        writer.writerow(feedback_data)
    
    return True

# Execute when user inputs data
if input_text:
    with st.spinner('Researching medication information...'):
        response = medicine_chain({"medicine": input_text})
    
    # Store generated content in session state
    st.session_state.generated_content = {
        "medicine": input_text,
        "info": response['medicine_info'],
        "side_effects": response['side_effects'],
        "precautions": response['precautions']
    }
    
    # Display results
    col1, col2 = st.columns(2)
    
    with col1:
        with st.expander('Medication Overview', expanded=True):
            st.info(response['medicine_info'])
    
    with col2:
        with st.expander('Potential Side Effects', expanded=True):
            st.warning(response['side_effects'])
    
    with st.expander('Precautions & Warnings', expanded=True):
        st.error(response['precautions'])
    
    # Chat interface section
    st.divider()
    st.subheader("💬 Ask Questions About " + input_text)
    
    # Clear chat history if medication changed
    if st.session_state.current_med != input_text:
        st.session_state.chat_history = []
        st.session_state.current_med = input_text
    
    # Create chat container
    chat_container = st.container(height=200, border=True)
    
    # Display chat history
    with chat_container:
        for message in st.session_state.chat_history:
            if message["role"] == "user":
                with st.chat_message("user"):
                    st.write(message["content"])
            elif message["role"] == "assistant":
                with st.chat_message("assistant"):
                    st.write(message["content"])
    
    # Question input
    user_question = st.chat_input(f"Ask about {input_text}...", key="chat_input")
    
    if user_question:
        # Add user question to chat history
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        
        # Display user question immediately
        with chat_container:
            with st.chat_message("user"):
                st.write(user_question)
        
        # Prepare context using session state
        context = f"""
        Medication Overview:
        {st.session_state.generated_content['info']}
        
        Side Effects:
        {st.session_state.generated_content['side_effects']}
        
        Precautions:
        {st.session_state.generated_content['precautions']}
        """
        
        # Create prompt for Q&A
        qa_prompt = PromptTemplate(
            input_variables=['context', 'question'],
            template="""
            You are a medical assistant providing information about medications.
            Answer the user's question based ONLY on the provided context.
            If the answer isn't in the context, politely state you don't know.
            Be concise but informative, and always include important safety information.
            
            Context:
            {context}
            
            Question: {question}
            Answer:
            """
        )
        
        # Create Q&A chain
        qa_chain = LLMChain(
            llm=ChatGroq(
                model_name="llama-3.3-70b-versatile", 
                temperature=0.2,
                api_key=os.environ['GROQ_API_KEY']
            ),
            prompt=qa_prompt
        )
        
        # Get response
        with st.spinner('Thinking...'):
            ai_response = qa_chain({
                'context': context,
                'question': user_question
            })
        
        # Add AI response to chat history
        st.session_state.chat_history.append({"role": "assistant", "content": ai_response['text']})
        
        # Display AI response
        with chat_container:
            with st.chat_message("assistant"):
                st.write(ai_response['text'])
        
        # Rerun to update chat display
        st.rerun()
    
    # Clear chat button
    if st.session_state.chat_history:
        if st.button("Clear Chat History"):
            st.session_state.chat_history = []
            st.rerun()
    
    # Quality evaluation section
    st.divider()
    st.subheader("Quality Evaluation")
    
    if st.button("📈 Evaluate Report Quality", help="Analyze the report for coherence, relevance, and accuracy"):
        metrics = evaluate_quality_metrics(
            input_text,
            response['medicine_info'],
            response['side_effects'],
            response['precautions']
        )
        st.session_state.quality_metrics = metrics
    
    if st.session_state.quality_metrics:
        display_quality_metrics(st.session_state.quality_metrics)
    
    # Report download section
    st.divider()
    st.subheader("Download Report")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # PDF Download
        try:
            pdf_report = create_pdf_report(
                input_text,
                response['medicine_info'],
                response['side_effects'],
                response['precautions']
            )
            st.download_button(
                label="📄 Download PDF Report",
                data=pdf_report,
                file_name=f"{input_text}_medication_report.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error(f"Error generating PDF: {str(e)}")
            st.warning("PDF generation failed. Please try the Word download instead.")
    
    with col2:
        # Word Download
        try:
            word_report = create_word_report(
                input_text,
                response['medicine_info'],
                response['side_effects'],
                response['precautions']
            )
            st.download_button(
                label="📝 Download Word Report",
                data=word_report,
                file_name=f"{input_text}_medication_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error generating Word document: {str(e)}")
    
    # User feedback section
    st.divider()
    st.subheader("User Feedback")
    
    if not st.session_state.feedback_submitted:
        with st.form("feedback_form"):
            st.write("How would you rate this report?")
            
            # Feedback type selection
            feedback_type = st.radio("Feedback type:", 
                                     ["General report", "Accuracy", "Completeness", "Usability"],
                                     horizontal=True)
            
            # Rating
            rating = st.slider("Overall rating:", 1, 5, 3,
                               help="1 = Poor, 5 = Excellent")
            
            # Comments
            comments = st.text_area("Specific feedback or suggestions:")
            
            # Submit button
            submitted = st.form_submit_button("Submit Feedback")
            
            if submitted:
                save_feedback(
                    input_text,
                    feedback_type,
                    rating,
                    comments
                )
                st.session_state.feedback_submitted = True
                st.success("Thank you for your feedback! It will help us improve our service.")
    else:
        st.info("Feedback submitted. Thank you for helping us improve!")
        if st.button("Submit new feedback"):
            st.session_state.feedback_submitted = False
            st.experimental_rerun()

# Feedback analysis section (for admin view)
if st.checkbox("Show feedback insights (Admin)"):
    try:
        feedback_df = pd.read_csv("user_feedback.csv")
        if not feedback_df.empty:
            st.subheader("Feedback Analysis")
            
            # Show raw data
            with st.expander("View Raw Feedback Data"):
                st.dataframe(feedback_df)
            
            # Show summary statistics
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Total Feedback Entries", len(feedback_df))
            with col2:
                avg_rating = feedback_df['rating'].mean()
                st.metric("Average Rating", f"{avg_rating:.1f}/5")
            
            # Show rating distribution
            st.subheader("Rating Distribution")
            rating_counts = feedback_df['rating'].value_counts().sort_index()
            st.bar_chart(rating_counts)
            
            # Show feedback types
            st.subheader("Feedback Types")
            type_counts = feedback_df['feedback_type'].value_counts()
            st.dataframe(type_counts)
        else:
            st.info("No feedback data available yet")
    except FileNotFoundError:
        st.warning("No feedback data collected yet")
    except Exception as e:
        st.error(f"Error loading feedback data: {str(e)}")

# Add footer
st.divider()
st.caption("Medication Information Assistant v1.0 | Powered by Groq and Llama 3")